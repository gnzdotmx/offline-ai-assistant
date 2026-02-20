"""
Tests for offline_ai_assistant.data.extractor.

Covers _clean_extracted_text, _calculate_file_hash, validate_file,
extract_from_file (with minimal PDF/DOCX when deps available), and extract_document.
"""

import stat
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch, MagicMock

from offline_ai_assistant.data.extractor import (
    DocumentExtractor,
    extract_document,
)


class TestCleanExtractedText(unittest.TestCase):
    def test_empty_returns_unchanged(self):
        self.assertEqual(DocumentExtractor._clean_extracted_text(""), "")
        self.assertEqual(DocumentExtractor._clean_extracted_text("   \n  "), "   \n  ")

    def test_merge_hyphenated_line_breaks(self):
        text = "word-\nnext"
        out = DocumentExtractor._clean_extracted_text(text)
        self.assertIn("word next", out)

    def test_remove_repeated_lines(self):
        text = "Header\nLine A\nLine B\nHeader\nLine C"
        out = DocumentExtractor._clean_extracted_text(text, window_size=5)
        # First "Header" kept, second duplicate removed
        lines = out.splitlines()
        self.assertEqual(lines.count("Header"), 1)
        self.assertIn("Line A", out)
        self.assertIn("Line C", out)

    def test_empty_lines_not_deduped_as_content(self):
        text = "A\n\n\nB"
        out = DocumentExtractor._clean_extracted_text(text)
        self.assertIn("A", out)
        self.assertIn("B", out)


class TestCalculateFileHash(unittest.TestCase):
    def test_hash_deterministic(self):
        extractor = DocumentExtractor()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as f:
            f.write(b"same content")
            path = Path(f.name)
        try:
            h1 = extractor._calculate_file_hash(path)
            h2 = extractor._calculate_file_hash(path)
            self.assertEqual(h1, h2)
            self.assertEqual(len(h1), 64)
            self.assertTrue(all(c in "0123456789abcdef" for c in h1))
        finally:
            path.unlink(missing_ok=True)

    def test_hash_different_for_different_content(self):
        extractor = DocumentExtractor()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as f:
            f.write(b"content A")
            path_a = Path(f.name)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as f:
            f.write(b"content B")
            path_b = Path(f.name)
        try:
            self.assertNotEqual(
                extractor._calculate_file_hash(path_a),
                extractor._calculate_file_hash(path_b),
            )
        finally:
            path_a.unlink(missing_ok=True)
            path_b.unlink(missing_ok=True)


class TestValidateFile(unittest.TestCase):
    def test_file_not_exists_returns_false(self):
        extractor = DocumentExtractor()
        ok, msg = extractor.validate_file(Path("/nonexistent/file.pdf"))
        self.assertFalse(ok)
        self.assertIn("exist", msg.lower())

    def test_directory_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp)
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                extractor = DocumentExtractor()
                ok, msg = extractor.validate_file(path)
                self.assertFalse(ok)
                self.assertIn("not a file", msg.lower())

    def test_unsupported_extension_returns_false(self):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            path = Path(f.name)
        try:
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                extractor = DocumentExtractor()
                ok, msg = extractor.validate_file(path)
                self.assertFalse(ok)
                self.assertIn("Unsupported", msg)
        finally:
            path.unlink(missing_ok=True)

    def test_file_too_large_returns_false(self):
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = Path(f.name)
        try:
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                # Patch Path.stat so validate_file sees a large file size
                big_stat = MagicMock(st_size=101 * 1024 * 1024, st_mode=stat.S_IFREG)
                with patch("pathlib.Path.stat", return_value=big_stat):
                    extractor = DocumentExtractor()
                    ok, msg = extractor.validate_file(path)
                    self.assertFalse(ok)
                    self.assertIn("large", msg.lower())
        finally:
            path.unlink(missing_ok=True)

    def test_valid_pdf_returns_true_when_fitz_available(self):
        try:
            import fitz
        except ImportError:
            fitz = None
        if fitz is None:
            self.skipTest("PyMuPDF not installed")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = Path(f.name)
        try:
            doc = fitz.open()
            doc.insert_page(0, width=100, height=100)
            doc[0].insert_text((10, 10), "test")
            doc.save(str(path))
            doc.close()
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                extractor = DocumentExtractor()
                ok, msg = extractor.validate_file(path)
                self.assertTrue(ok, msg)
                self.assertEqual(msg, "")
        finally:
            path.unlink(missing_ok=True)


class TestExtractFromFile(unittest.TestCase):
    def test_file_not_found_raises(self):
        extractor = DocumentExtractor()
        with self.assertRaises(FileNotFoundError) as ctx:
            extractor.extract_from_file(Path("/nonexistent/file.pdf"))
        self.assertIn("not found", str(ctx.exception).lower())

    def test_unsupported_type_raises(self):
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            path = Path(f.name)
        try:
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                extractor = DocumentExtractor()
                with self.assertRaises(ValueError) as ctx:
                    extractor.extract_from_file(path)
                self.assertIn("Unsupported", str(ctx.exception))
        finally:
            path.unlink(missing_ok=True)

    def test_extract_pdf_returns_expected_keys(self):
        try:
            import fitz
        except ImportError:
            fitz = None
        if fitz is None:
            self.skipTest("PyMuPDF not installed")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = Path(f.name)
        try:
            doc = fitz.open()
            doc.insert_page(0, width=200, height=200)
            doc[0].insert_text((20, 20), "Hello PDF")
            doc.save(str(path))
            doc.close()
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                cfg.EXTRACTOR_CLEAN_TEXT = False
                extractor = DocumentExtractor()
                result = extractor.extract_from_file(path)
            self.assertIn("full_text", result)
            self.assertIn("file_hash", result)
            self.assertIn("file_name", result)
            self.assertIn("file_type", result)
            self.assertEqual(result["file_type"], "pdf")
            self.assertIn("pages", result)
            self.assertGreaterEqual(len(result["pages"]), 1)
            self.assertIn("Hello", result["full_text"] or "")
        finally:
            path.unlink(missing_ok=True)

    def test_extract_docx_returns_expected_keys(self):
        try:
            from docx import Document
        except ImportError:
            Document = None
        if Document is None:
            self.skipTest("python-docx not installed")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        try:
            doc = Document()
            doc.add_paragraph("Hello DOCX")
            doc.save(path)
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                cfg.EXTRACTOR_CLEAN_TEXT = False
                extractor = DocumentExtractor()
                result = extractor.extract_from_file(path)
            self.assertIn("full_text", result)
            self.assertIn("file_hash", result)
            self.assertEqual(result["file_type"], "docx")
            self.assertIn("paragraphs", result)
            self.assertIn("Hello", result["full_text"] or "")
        finally:
            path.unlink(missing_ok=True)


class TestGetSupportedExtensions(unittest.TestCase):
    def test_returns_list(self):
        extractor = DocumentExtractor()
        exts = extractor.get_supported_extensions()
        self.assertIsInstance(exts, list)
        self.assertTrue(all(e.startswith(".") for e in exts))

    def test_contains_only_pdf_or_docx(self):
        extractor = DocumentExtractor()
        exts = extractor.get_supported_extensions()
        for e in exts:
            self.assertIn(e, {".pdf", ".docx"})


class TestExtractDocument(unittest.TestCase):
    def test_extract_document_returns_dict(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        try:
            doc = Document()
            doc.add_paragraph("Test")
            doc.save(path)
            with patch("offline_ai_assistant.data.extractor.Config") as cfg:
                cfg.SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
                cfg.EXTRACTOR_CLEAN_TEXT = False
                result = extract_document(path)
                self.assertIsInstance(result, dict)
                self.assertIn("full_text", result)
        finally:
            path.unlink(missing_ok=True)
